import docx
from PIL import Image
from docx.shared import RGBColor
import random
import string
import os
import sys
from nltk.corpus import words
import random
import string
import nltk
nltk.download('treebank')
from docx.shared import Cm, Pt

# create a new Word file
doc = docx.Document()
image_path = sys.argv[1]
def get_rgb_matrix(image_path):
    # open image file
    image = Image.open(image_path)

    # get size of image
    width, height = image.size

    # create matrix of RGB values
    rgb_matrix = [[0 for x in range(width)] for y in range(height)]

    # iterate over each pixel in the image
    for x in range(width):
        for y in range(height):
            # get RGB values of pixel
            r, g, b = image.getpixel((x, y))
            # assign RGB values to matrix
            rgb_matrix[y][x] = (r, g, b)

    return rgb_matrix

from lorem_text import lorem

# def generate_random_text(n):
#     """
#     Generates random text of n words.
#     """
#     return lorem.words(n)

# def generate_random_string(n):
#     # Define the character set for the random string
#     characters = string.ascii_letters + string.digits

#     # Generate the random string
#     random_string = ''.join(random.choices(characters, k=n))

#     return random_string

import string
import random

def generate_string(n):
    letters = string.ascii_lowercase
    result_str = ''.join(random.choice(letters) for i in range(n))
    return result_str

# n = 10  # length of the string to generate
# generated_string = generate_string(n)
# print(generated_string)


# def append_word(document, rgb, text):
#     color = RGBColor(rgb[0], rgb[1], rgb[2])
#     # Get the last paragraph in the document
#     last_paragraph = document.paragraphs[-1]

#     # Add a run with the specified text and font color
#     run = last_paragraph.add_run(text)
#     font = run.font
#     font.color.rgb = RGBColor(*color)
    
def append_word(document, rgb, text):
    # Check if the document is empty
    color = RGBColor(rgb[0], rgb[1], rgb[2])
    
    if len(document.paragraphs) == 0:
        # If the document is empty, add a new paragraph
        document.add_paragraph()

    # Get the last paragraph in the document
    last_paragraph = document.paragraphs[-1]

    # Add a run with the specified text and font color
    run = last_paragraph.add_run(text)
    font = run.font
    font.color.rgb = RGBColor(*color)
    
def is_word_file_empty():
    # Open the document
    

    # Iterate through all the paragraphs in the document
    for paragraph in doc.paragraphs:
        # Check if the paragraph contains any text
        if paragraph.text:
            return False

    # If no non-empty paragraphs were found, the document is empty
    return True

# def reduce_image_size(image_path, new_width):
#     # Open the image file
#     image = Image.open(image_path)

#     # Calculate the new height based on the new width while maintaining aspect ratio
#     width, height = image.size
#     new_height = int(height * (new_width / width))

#     # Resize the image to the new dimensions
#     resized_image = image.resize((new_width, new_height))

#     # Save the resized image to a new file
#     resized_image.save('resized_image.jpg')


def generate_sentence(n):
    words = []
    while sum(len(word) for word in words) + len(words) - 1 < n:
        word = ''.join(random.choice(string.ascii_lowercase) for _ in range(random.randint(3, 8)))
        words.append(word)
    return ' '.join(words).capitalize() + '.'



def steganography(rgb):
    (x, y) = (len(rgb), len(rgb[0]))
    sentence_list = nltk.corpus.treebank.sents()
    text = ''
    # text = generate_sentence(x*y)

    # Generate random meaningful English sentences and concatenate them into a paragraph
    for i in range(x*y):
        sentence = ' '.join(random.choice(sentence_list))
        text += sentence + ' '
    doc2 = docx.Document()

    # add some text to the document
    doc2.add_paragraph(text)

    # save the document
    # doc2.save('my_document.docx')
    if (os.path.exists(sys.argv[2])):
        os.remove(sys.argv[2])
        doc2.save(sys.argv[2]) 

    else :
        doc2.save(sys.argv[2]) 
    
    text_it = 0
    
    for i in range(x):
        # if ()
        # last_paragraph = doc.paragraphs[-1]
        for j in range(y):
            # print(i,j)
            # add a paragraph to the file
            rgb_val = rgb[i][j]
            
            word = text[text_it]
            text_it+=1
            append_word(doc,rgb_val,word)
        doc.add_paragraph('\n')


rgb = get_rgb_matrix(image_path)
# print(rgb[0][0])
# reduce_image_size(image_path, int(len(rgb[0])/3))
# rgb = get_rgb_matrix('resized_image.jpg')
# # print(len(rgb[0]), len(rgb))
print("Please wait...")
steganography(rgb)
doc.styles['Normal'].font.size = Pt(1)
# set the page size to 54x54cm
section = doc.sections[0]
section.page_width = Cm(54)
section.page_height = Cm(54)
for paragraph in doc.paragraphs:
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)

if (os.path.exists(sys.argv[3])):
    os.remove(sys.argv[3])
    doc.save(sys.argv[3]) 

else :
    doc.save(sys.argv[3])        

print("Image Hidden Successfully.")
            
